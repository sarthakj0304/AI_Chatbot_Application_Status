import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def create_docx(filename, title, content):
    doc = Document()
    doc.add_heading(title, 0)
    for paragraph in content.split('\n\n'):
        doc.add_paragraph(paragraph.strip())
    doc.save(filename)

def create_pdf(filename, title, content):
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, height - 72, title)
    
    c.setFont("Helvetica", 12)
    y_position = height - 100
    
    for line in content.split('\n'):
        if y_position < 72:
            c.showPage()
            c.setFont("Helvetica", 12)
            y_position = height - 72
        if line.strip():
            c.drawString(72, y_position, line.strip())
            y_position -= 15
            
    c.save()

def main():
    base_dir = "/Users/sarthakjain/Desktop/AI_Chatbot_Application_Status-main/test_docs"
    
    # Document 1: Engineering Guidelines (DOCX)
    eng_title = "Acme Corp Engineering Guidelines 2026"
    eng_content = """
Welcome to the Acme Corp Engineering team!

Code Review Process:
All pull requests must be approved by at least two senior engineers before merging to the main branch. Code reviews should be completed within 24 hours.

Deployment Schedule:
Production deployments occur every Tuesday and Thursday at 10:00 AM PST. Hotfixes can be deployed at any time with VP approval.

On-Call Rotation:
Engineers are expected to participate in the on-call rotation. Each rotation lasts for one week, starting on Monday at 9:00 AM. On-call engineers receive an additional $500 stipend per week.
"""
    create_docx(os.path.join(base_dir, "engineering_guidelines.docx"), eng_title, eng_content)
    
    # Document 2: Employee Benefits (PDF)
    benefits_title = "Acme Corp Employee Benefits & Perks"
    benefits_content = """
Acme Corp values our employees and offers a comprehensive benefits package.

Health Insurance:
We offer full medical, dental, and vision coverage for all full-time employees and their dependents. Acme Corp covers 100% of the premiums.

Paid Time Off (PTO):
Employees receive 20 days of paid time off per year, which accrues starting from day one. Additionally, we observe 11 standard company holidays.

Retirement Plan:
Acme Corp offers a 401(k) retirement plan with a 4% company match. Employees are fully vested immediately upon enrollment.

Fitness Stipend:
To promote health and wellness, employees receive a $100 monthly stipend for gym memberships or fitness classes.
"""
    create_pdf(os.path.join(base_dir, "employee_benefits.pdf"), benefits_title, benefits_content)
    
    print("Test documents generated successfully in 'test_docs' directory.")

if __name__ == "__main__":
    main()
